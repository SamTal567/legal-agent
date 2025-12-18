import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Setup paths
base_dir = os.path.join("legal_agent", "data", "templates")
os.makedirs(base_dir, exist_ok=True)

def create_legal_notice():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Header
    head = doc.add_paragraph("ADVOCATE AI LEGAL SERVICES")
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.runs[0].bold = True
    head.runs[0].font.size = Pt(16)
    
    doc.add_paragraph("Office: Cloud Server, Internet | Email: advocate.ai@legalagent.com")
    doc.add_paragraph("-" * 90)

    # Ref and Date
    p = doc.add_paragraph()
    p.add_run("Ref No: LEGAL/2025/{{REF_NO}}\t\t\t\t\t")
    p.add_run("Date: {{DATE}}").bold = True

    p = doc.add_paragraph("REGISTERED POST WITH A/D / SPEED POST")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].underline = True

    # Recipient
    doc.add_paragraph("To,\n{{OPPONENT_NAME}}\n{{OPPONENT_ADDRESS}}")

    # Subject
    p = doc.add_paragraph()
    p.add_run("SUBJECT: LEGAL NOTICE FOR {{REASON}}").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Body
    doc.add_paragraph("Sir/Madam,")
    
    p = doc.add_paragraph("Under instructions from and on behalf of my client, ")
    p.add_run("{{CLIENT_NAME}}").bold = True
    p.add_run(", resident of {{CLIENT_ADDRESS}}, I hereby serve you with the following legal notice:")

    doc.add_paragraph("1. That my client states: {{CASE_DETAILS}}")
    doc.add_paragraph("2. That despite repeated requests, you have failed to fulfill your obligations.")
    doc.add_paragraph("3. That your act has caused mental agony and financial loss to my client.")

    # Demand
    p = doc.add_paragraph()
    p.add_run("I THEREFORE CALL UPON YOU").bold = True
    p.add_run(" to {{DEMAND}} within 15 days from the receipt of this notice, failing which my client shall be constrained to initiate appropriate civil/criminal proceedings against you at your own risk and cost.")

    doc.add_paragraph("You are also liable to pay Rs. 5,000/- as the cost of this legal notice.")
    
    doc.add_paragraph("A copy of this notice has been retained in my office for record and further necessary action.")

    doc.add_paragraph("\nSincerely,\n")
    doc.add_paragraph("ADVOCATE AI")
    
    path = os.path.join(base_dir, "LEGAL NOTICE To.docx")
    doc.save(path)
    print(f"Created: {path}")

def create_consumer_complaint():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    
    # Title
    head = doc.add_paragraph("BEFORE THE DISTRICT CONSUMER DISPUTES REDRESSAL COMMISSION AT {{CITY}}")
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.runs[0].bold = True
    head.runs[0].font.size = Pt(14)

    doc.add_paragraph("\nConsumer Complaint No. _______ of 2025")

    # Parties
    doc.add_paragraph("\nIN THE MATTER OF:")
    p = doc.add_paragraph("{{CLIENT_NAME}}\nR/o {{CLIENT_ADDRESS}}")
    p.runs[0].bold = True
    doc.add_paragraph("... COMPLAINANT")
    
    doc.add_paragraph("\nVERSUS\n")
    
    p = doc.add_paragraph("{{OPPONENT_NAME}}\nAddress: {{OPPONENT_ADDRESS}}")
    p.runs[0].bold = True
    doc.add_paragraph("... OPPOSITE PARTY")

    # Complaint Body
    doc.add_paragraph("\nCOMPLAINT UNDER SECTION 35 OF THE CONSUMER PROTECTION ACT, 2019")
    
    doc.add_paragraph("MOST RESPECTFULLY SHOWETH:")
    doc.add_paragraph("1. That the Complainant is a consumer as defined under the Act.")
    doc.add_paragraph("2. That the Complainant purchased/availed service: {{CASE_DETAILS}}")
    doc.add_paragraph("3. That the Opposite Party provided defective goods/deficient services by: {{DEFECT_DETAILS}}")
    doc.add_paragraph("4. That the Complainant approached the Opposite Party multiple times but received no relief.")
    
    # Jurisdiction
    doc.add_paragraph("5. JURISDICTION: The cause of action arose at {{CITY}}, and the value of goods/services is within the pecuniary limits of this Hon'ble Commission.")
    doc.add_paragraph("6. LIMITATION: The complaint is filed within the statutory period of 2 years.")

    # Prayer
    doc.add_paragraph("\nPRAYER:")
    doc.add_paragraph("It is therefore most respectfully prayed that this Hon'ble Commission may be pleased to:")
    doc.add_paragraph("a) Direct the Opposite Party to refund/replace/compensate as per law.")
    doc.add_paragraph("b) Pay Rs. {{COMPENSATION_AMOUNT}} as compensation for mental harassment.")
    doc.add_paragraph("c) Pay Rs. 10,000/- towards litigation costs.")

    doc.add_paragraph("\nCOMPLAINANT")
    doc.add_paragraph("Through Advocate AI")
    
    # Verification
    doc.add_paragraph("\nVERIFICATION")
    doc.add_paragraph("I, the complainant above, do hereby verify that the contents of paras 1 to 6 are true to my personal knowledge. Verified at {{CITY}} on {{DATE}}.")
    doc.add_paragraph("\nDEPONENT")

    path = os.path.join(base_dir, "Consumer_Complaint.docx")
    doc.save(path)
    print(f"Created: {path}")

def create_rti_app():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'

    doc.add_paragraph("APPLICATION FOR INFORMATION UNDER RIGHT TO INFORMATION ACT, 2005").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\nTo,\nThe Public Information Officer (PIO),\n{{DEPARTMENT_NAME}}\n{{DEPARTMENT_ADDRESS}}")
    
    doc.add_paragraph("\n1. Name of Applicant: {{CLIENT_NAME}}")
    doc.add_paragraph("2. Address: {{CLIENT_ADDRESS}}")
    
    doc.add_paragraph("\n3. Particulars of Information Sought:")
    doc.add_paragraph("   Subject: {{SUBJECT}}")
    doc.add_paragraph("   Period: {{PERIOD}}")
    doc.add_paragraph("   Details: {{CASE_DETAILS}}")
    
    doc.add_paragraph("\n4. I state that the information sought does not fall within the exemptions contained in Section 8 of the RTI Act.")
    doc.add_paragraph("5. A fee of Rs. 10/- has been deposited via Postal Order/Demand Draft.")
    
    doc.add_paragraph("\nPlace: {{CITY}}")
    doc.add_paragraph("Date: {{DATE}}")
    
    doc.add_paragraph("\nSignature of Applicant")

    path = os.path.join(base_dir, "RTI_Application.docx")
    doc.save(path)
    print(f"Created: {path}")

if __name__ == "__main__":
    create_legal_notice()
    create_consumer_complaint()
    create_rti_app()
    print("All professional templates created successfully.")