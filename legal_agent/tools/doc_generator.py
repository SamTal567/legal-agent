import os
import json
from docx import Document
from datetime import datetime

# Define where templates live and where drafts go
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_DIR = os.path.join(BASE_DIR, "data", "templates")
OUTPUT_DIR = os.path.join(BASE_DIR, "output")

if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

def replace_text_in_paragraph(paragraph, key, value):
    """
    Replaces text while trying to preserve formatting.
    """
    if key in paragraph.text:
        # Strategy 1: Check individual runs (Preserves Bold/Italic)
        inline = paragraph.runs
        replaced = False
        for i in range(len(inline)):
            if key in inline[i].text:
                inline[i].text = inline[i].text.replace(key, str(value))
                replaced = True
        
        # Strategy 2: Fallback to paragraph replacement
        if not replaced:
            paragraph.text = paragraph.text.replace(key, str(value))

def generate_legal_document(doc_type: str, user_details_json: str) -> str:
    """
    Generate a legal document by filling template with user details.
    """
    try:
        # 1. Parse Data & Normalize Keys
        raw_data = json.loads(user_details_json)
        data = {}
        
        # KEY FIX: Convert everything to UPPERCASE keys to match Template {{KEY}} style
        # Also map common AI mistakes to Template Keys
        key_mapping = {
            "DETAILS": "CASE_DETAILS",
            "CASE": "CASE_DETAILS",
            "REF": "REF_NO",
            "REFERENCE": "REF_NO",
            "OPPONENT": "OPPONENT_NAME",
            "CLIENT": "CLIENT_NAME",
            "DEFECT": "DEFECT_DETAILS"
        }

        for k, v in raw_data.items():
            upper_k = k.upper().replace(" ", "_") # "Ref No" -> "REF_NO"
            
            # Use mapped key if available, else use the uppercase key
            final_key = key_mapping.get(upper_k, upper_k)
            data[final_key] = v

        # 2. Find Template
        candidates = [
            f"{doc_type}.docx",
            f"LEGAL NOTICE To.docx" if "notice" in doc_type.lower() else None,
            "Consumer_Complaint.docx" if "consumer" in doc_type.lower() else None,
            "RTI_Application.docx" if "rti" in doc_type.lower() else None
        ]
        
        template_path = None
        for name in candidates:
            if name:
                path = os.path.join(TEMPLATE_DIR, name)
                if os.path.exists(path):
                    template_path = path
                    break
        
        if not template_path:
            return f"Error: No template found for '{doc_type}' in {TEMPLATE_DIR}"

        # 3. Load Document
        doc = Document(template_path)
        
        # 4. Replace Placeholders
        for key, value in data.items():
            # Force {{KEY}} format
            placeholder = "{{" + key + "}}"
            
            # Replace in Paragraphs
            for paragraph in doc.paragraphs:
                replace_text_in_paragraph(paragraph, placeholder, value)

            # Replace in Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, placeholder, value)

        # 5. Save
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Draft_{doc_type}_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_DIR, filename)
        
        doc.save(output_path)
        
        return f"Success! Document created at: {output_path}"

    except Exception as e:
        return f"Failed to generate document: {str(e)}"